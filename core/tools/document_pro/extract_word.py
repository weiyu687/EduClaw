"""
提取 word 内容
功能: 支持 doc/docx，文本与表格按物理顺序混合排列
Author: Gongmin Wei
Date: 2026-04-05
"""
import json
import io
from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
import aspose.words as aw


def iter_block_items(parent):
    """
    遍历 Word 文档中的所有块级元素(段落 + 表格)，保持物理顺序
    """
    if isinstance(parent, DocumentObject):
        parent_elm = parent.element.body
    else:
        raise ValueError("Unsupported parent type")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def extract_word(word_path: str) -> str:
    """
    从 Word 文件中提取内容 (文本 + 表格)，支持 doc/docx
    """
    result = {
        "metadata": {
            "total_paragraphs": 0,
            "total_tables": 0,
            "total_characters": 0
        },
        "elements": []
    }

    try:
        if word_path.lower().endswith(".doc"):
            doc_aw = aw.Document(word_path)
            temp_stream = io.BytesIO()
            doc_aw.save(temp_stream, aw.SaveFormat.DOCX)
            doc = Document(temp_stream)
        else:
            doc = Document(word_path)

        for block in iter_block_items(doc):
            if isinstance(block, Paragraph):
                text = block.text.strip()
                if text:
                    result["elements"].append({"type": "text", "content": text})
                    result["metadata"]["total_paragraphs"] += 1
                    result["metadata"]["total_characters"] += len(text)

            elif isinstance(block, Table):
                rows = []
                for row in block.rows:
                    cells = [cell.text.strip() if cell.text else "" for cell in row.cells]
                    rows.append(cells)

                if rows:
                    md_table = []
                    for i, row in enumerate(rows):
                        line = "| " + " | ".join(row) + " |"
                        md_table.append(line)
                        if i == 0:
                            md_table.append("|" + "|".join(["---"] * len(row)) + "|")

                    table_md = "\n".join(md_table)
                    result["elements"].append({"type": "table", "content": table_md})
                    result["metadata"]["total_tables"] += 1

    except Exception as e:
        return json.dumps({"error": f"Word 文件读取失败: {str(e)}"}, ensure_ascii=False, indent=2)

    return json.dumps(result, ensure_ascii=False, indent=2)